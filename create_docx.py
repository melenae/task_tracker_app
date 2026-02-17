#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для конвертации Markdown в DOCX
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_break(doc):
    """Добавить разрыв страницы"""
    doc.add_page_break()

def markdown_to_docx(md_file, docx_file):
    """Конвертирует Markdown файл в DOCX"""
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаем новый документ Word
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Обработка кодовых блоков
        if stripped.startswith('```'):
            if in_code_block:
                # Закрываем блок кода
                if code_block_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    # Серый фон для кода
                    p.style = 'No Spacing'
                code_block_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Открываем блок кода
                code_language = stripped[3:].strip()
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Пропускаем пустые строки (но не все подряд)
        if not stripped:
            if i + 1 < len(lines) and lines[i + 1].strip():
                doc.add_paragraph()
            i += 1
            continue
        
        # Горизонтальная линия
        if stripped.startswith('---') or stripped.startswith('==='):
            p = doc.add_paragraph('─' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Заголовки
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('##### '):
            p = doc.add_heading(stripped[6:], level=5)
        elif stripped.startswith('###### '):
            p = doc.add_heading(stripped[7:], level=6)
        
        # Таблицы
        elif '|' in stripped:
            # Проверяем, это разделитель таблицы?
            if re.match(r'^\|[\s\-\|:]+\|$', stripped):
                # Это разделитель, пропускаем
                i += 1
                continue
            
            # Собираем данные таблицы
            if not in_table:
                in_table = True
                table_data = []
            
            # Парсим строку таблицы
            cells = [cell.strip() for cell in stripped.split('|')]
            # Убираем пустые элементы в начале и конце
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            
            if cells:
                table_data.append(cells)
        
        # Нумерованные списки
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            # Обработка форматирования в тексте
            p = doc.add_paragraph(text, style='List Number')
            _format_text(p, text)
        
        # Маркированные списки
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _format_text(p, text)
        
        # Обычный текст
        else:
            if in_table:
                # Завершаем таблицу
                if table_data:
                    _create_table(doc, table_data)
                table_data = []
                in_table = False
            
            # Обработка форматирования
            p = doc.add_paragraph()
            _format_text(p, stripped)
        
        i += 1
    
    # Если таблица не была закрыта
    if in_table and table_data:
        _create_table(doc, table_data)
    
    # Сохраняем документ
    doc.save(docx_file)
    print(f"✅ Документ создан: {docx_file}")

def _format_text(paragraph, text):
    """Форматирует текст с поддержкой Markdown"""
    # Паттерны для форматирования
    patterns = [
        (r'\*\*(.*?)\*\*', True, False),  # Жирный
        (r'\*(.*?)\*', False, True),       # Курсив
        (r'`(.*?)`', False, False, 'code'), # Код
        (r'~~(.*?)~~', False, False, 'strike'), # Зачеркнутый
    ]
    
    # Разделяем текст на части
    parts = []
    last_pos = 0
    
    # Находим все совпадения
    matches = []
    for pattern, bold, italic, *extra in patterns:
        for match in re.finditer(pattern, text):
            matches.append((match.start(), match.end(), match.group(1), bold, italic, extra[0] if extra else None))
    
    # Сортируем по позиции
    matches.sort(key=lambda x: x[0])
    
    # Создаем части
    for start, end, content, bold, italic, extra in matches:
        if last_pos < start:
            parts.append(('text', text[last_pos:start]))
        parts.append(('formatted', content, bold, italic, extra))
        last_pos = end
    
    if last_pos < len(text):
        parts.append(('text', text[last_pos:]))
    
    # Добавляем текст в параграф
    if not parts:
        paragraph.add_run(text)
    else:
        for part_type, *args in parts:
            if part_type == 'text':
                paragraph.add_run(args[0])
            else:
                content, bold, italic, extra = args
                run = paragraph.add_run(content)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if extra == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                elif extra == 'strike':
                    run.font.strike = True

def _create_table(doc, table_data):
    """Создает таблицу в документе"""
    if not table_data:
        return
    
    # Определяем количество колонок
    max_cols = max(len(row) for row in table_data) if table_data else 1
    
    # Создаем таблицу
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    # Заполняем таблицу
    for row_idx, row_data in enumerate(table_data):
        for col_idx in range(max_cols):
            cell = table.rows[row_idx].cells[col_idx]
            if col_idx < len(row_data):
                cell.text = row_data[col_idx]
                # Убираем форматирование Markdown из ячеек
                cell.text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell.text)
                cell.text = re.sub(r'\*(.*?)\*', r'\1', cell.text)
                cell.text = re.sub(r'`(.*?)`', r'\1', cell.text)
            else:
                cell.text = ''
    
    # Делаем первую строку заголовком (если есть)
    if len(table_data) > 1:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

if __name__ == '__main__':
    import sys
    
    md_file = 'ПРОЕКТ_ОПИСАНИЕ.md'
    docx_file = 'ПРОЕКТ_ОПИСАНИЕ.docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    
    try:
        markdown_to_docx(md_file, docx_file)
        print(f"✅ Успешно создан файл: {docx_file}")
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        import traceback
        traceback.print_exc()














